"""
Vector Database Loader - Handles content extraction, embedding creation, and storage
"""
from pathlib import Path
from typing import List, Dict, Optional
import logging
import json
import requests
from datetime import datetime
from agenticaicapstone.src.utils.azure_openai_client import AzureOpenAIClient

logger = logging.getLogger(__name__)


class VectorDBLoader:
    """Handles loading invoice documents into vector database"""
    
    def __init__(self, vector_store=None, openai_client=None, api_base_url="http://localhost:8000"):
        """
        Initialize the vector database loader
        
        Args:
            vector_store: Optional InvoiceVectorStore instance. If None, will be lazily loaded.
            openai_client: Optional AzureOpenAIClient instance. If None, will be lazily loaded.
            api_base_url: Base URL for API endpoints (default: http://localhost:8000)
        """
        self._vector_store = vector_store
        self._openai_client = openai_client
        self.api_base_url = api_base_url
    
    @property
    def vector_store(self):
        """Lazy load vector store when needed"""
        if self._vector_store is None:
            try:
                from agenticaicapstone.src.rag.vector_store import InvoiceVectorStore
                self._vector_store = InvoiceVectorStore()
            except Exception as e:
                logger.error(f"Failed to initialize vector store: {e}")
                raise
        return self._vector_store
    
    @property
    def openai_client(self):
        """Lazy load Azure OpenAI client when needed"""
        if self._openai_client is None:
            try:
                self._openai_client = AzureOpenAIClient()
            except Exception as e:
                logger.error(f"Failed to initialize Azure OpenAI client: {e}")
                raise
        return self._openai_client
    
    def _get_next_invoice_id(self) -> str:
        """Get next available invoice ID from API"""
        try:
            response = requests.get(f"{self.api_base_url}/vector/next-id/invoice", timeout=60)
            response.raise_for_status()
            return response.json()['id']
        except Exception as e:
            logger.error(f"Failed to get next invoice ID from API: {e}")
            # Fallback to timestamp-based ID
            return f"INV-{datetime.now().strftime('%Y%m%d%H%M%S')}"
    
    def _get_next_vendor_ids(self, count: int) -> List[str]:
        """Get next N vendor IDs from API (calls single-ID endpoint N times)"""
        try:
            vendor_ids = []
            for _ in range(count):
                response = requests.get(f"{self.api_base_url}/vector/next-id/vendor", timeout=60)
                response.raise_for_status()
                vendor_ids.append(response.json()['id'])
            return vendor_ids
        except Exception as e:
            logger.error(f"Failed to get vendor IDs from API: {e}")
            # Fallback to timestamp-based IDs
            base_time = datetime.now()
            return [f"VND-{base_time.strftime('%Y%m%d%H%M%S')}{i:02d}" for i in range(count)]
    
    def _get_next_sku_ids(self, count: int) -> List[str]:
        """Get next N SKU IDs from API (calls single-ID endpoint N times)"""
        try:
            sku_ids = []
            for _ in range(count):
                response = requests.get(f"{self.api_base_url}/vector/next-id/sku", timeout=60)
                response.raise_for_status()
                sku_ids.append(response.json()['id'])
            return sku_ids
        except Exception as e:
            logger.error(f"Failed to get SKU IDs from API: {e}")
            # Fallback to timestamp-based IDs
            base_time = datetime.now()
            return [f"SKU-{base_time.strftime('%Y%m%d%H%M%S')}{i:02d}" for i in range(count)]
    
    def extract_content(self, file_path: Path) -> str:
        """
        Extract text content from a file based on its type
        
        Args:
            file_path: Path to the file
            
        Returns:
            str: Extracted text content
            
        Raises:
            ValueError: If file type is not supported
        """
        file_ext = file_path.suffix.lower()
        
        try:
            if file_ext == '.json':
                return self._extract_json(file_path)
            elif file_ext == '.csv':
                return self._extract_csv(file_path)
            elif file_ext == '.pdf':
                return self._extract_pdf(file_path)
            elif file_ext == '.docx':
                return self._extract_docx(file_path)
            elif file_ext in ['.jpg', '.png']:
                return self._extract_image(file_path)
            else:
                raise ValueError(f"Unsupported file type: {file_ext}")
                
        except Exception as e:
            logger.error(f"Failed to extract content from {file_path}: {e}")
            raise
    
    def _extract_json(self, file_path: Path) -> str:
        """Extract content from JSON file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            # Convert JSON to readable text
            return json.dumps(data, ensure_ascii=False)
        except Exception as e:
            raise ValueError(f"Failed to parse JSON file: {e}")
    
    def _extract_csv(self, file_path: Path) -> str:
        """Extract content from CSV file"""
        try:
            import csv
            with open(file_path, 'r', encoding='utf-8') as f:
                reader = csv.reader(f)
                rows = list(reader)
            # Convert CSV to readable text
            return '\n'.join([','.join(row) for row in rows])
        except Exception as e:
            raise ValueError(f"Failed to parse CSV file: {e}")
    
    def _analyze_image_with_vision(self, image_base64: str, context: str = "") -> str:
        """Analyze image using Azure OpenAI Vision API
        
        Args:
            image_base64: Base64 encoded image data
            context: Optional context about where the image came from (e.g., 'page 1', 'document')
            
        Returns:
            str: Analysis result or error message
        """
        try:
            prompt = (
                "Analyze this image from an invoice/document. "
                "Extract all visible text, numbers, tables, stamps, signatures, "
                "and any other relevant information. Be detailed and structured. "
                "IMPORTANT: Provide your response in English only, even if the document is in another language. "
                "Translate any non-English text to English EXCEPT any _id field, keep them as is."
            )
            
            analysis = self.openai_client.chat_with_image(
                prompt=prompt,
                image_base64=image_base64,
                temperature=0.3,
                max_tokens=1000
            )
            
            return analysis
            
        except Exception as e:
            logger.error(f"Failed to analyze image {context}: {e}")
            return f"[Failed to analyze: {e}]"
    
    def _extract_pdf(self, file_path: Path) -> str:
        """Extract text from PDF file"""
        try:
            import PyPDF2
        except ImportError:
            logger.error("PyPDF2 not installed. Install with: pip install PyPDF2")
            raise ValueError("PyPDF2 library not available")
        
        try:
            import base64
            import io
            from PIL import Image
        except ImportError:
            logger.warning("PIL not installed. Image extraction from PDF will be disabled.")
            Image = None
        
        try:
            text_content = []
            image_analyses = []
            
            with open(file_path, 'rb') as pdf_file:
                pdf_reader = PyPDF2.PdfReader(pdf_file)
                
                # Get number of pages
                num_pages = len(pdf_reader.pages)
                
                # Extract text and images from each page
                for page_num in range(num_pages):
                    page = pdf_reader.pages[page_num]
                    page_text = page.extract_text()
                    
                    if page_text.strip():  # Only add non-empty pages
                        text_content.append(f"--- Page {page_num + 1} ---\n{page_text}")
                    
                    # Extract images from page if PIL is available
                    if Image is not None and '/XObject' in page['/Resources']:
                        xObject = page['/Resources']['/XObject'].get_object()
                        
                        for obj in xObject:
                            if xObject[obj]['/Subtype'] == '/Image':
                                try:
                                    # Get image data
                                    img_obj = xObject[obj]
                                    img_data = img_obj.get_data()
                                    
                                    # Convert to base64
                                    img_base64 = base64.b64encode(img_data).decode('utf-8')
                                    
                                    # Analyze image using common method
                                    analysis = self._analyze_image_with_vision(
                                        img_base64, 
                                        context=f"from page {page_num + 1}"
                                    )
                                    
                                    image_analyses.append(
                                        f"--- Image from Page {page_num + 1} ---\n{analysis}"
                                    )
                                        
                                except Exception as e:
                                    logger.warning(f"Failed to extract image from page {page_num + 1}: {e}")
                
                # Combine text and image analyses
                all_content = []
                
                if text_content:
                    all_content.append("=== PDF TEXT CONTENT ===\n" + "\n\n".join(text_content))
                
                if image_analyses:
                    all_content.append("=== IMAGE ANALYSIS ===\n" + "\n\n".join(image_analyses))
                
                extracted_text = "\n\n".join(all_content)
                
                if not extracted_text.strip():
                    logger.warning(f"No text or images extracted from PDF: {file_path.name}")
                    return f"[No content found in {file_path.name}]"
                
                return extracted_text
                
        except PyPDF2.errors.PdfReadError as e:
            logger.error(f"PDF read error for {file_path.name}: {e}")
            raise ValueError(f"Failed to read PDF file: {e}")
        except Exception as e:
            logger.error(f"Unexpected error extracting PDF {file_path.name}: {e}")
            raise ValueError(f"Failed to extract PDF: {e}")
    
    def _extract_docx(self, file_path: Path) -> str:
        """Extract text and images from DOCX file"""
        try:
            from docx import Document
        except ImportError:
            logger.error("python-docx not installed. Install with: pip install python-docx")
            raise ValueError("python-docx library not available")
        
        try:
            import base64
            from PIL import Image as PILImage
            import io
        except ImportError:
            logger.warning("PIL not installed. Image extraction from DOCX will be disabled.")
            PILImage = None
        
        try:
            text_content = []
            image_analyses = []
            
            # Load the document
            doc = Document(file_path)
            
            logger.info(f"Extracting content from {len(doc.paragraphs)} paragraphs in {file_path.name}")
            
            # Extract text from paragraphs
            for i, para in enumerate(doc.paragraphs):
                text = para.text.strip()
                if text:
                    text_content.append(text)
            
            # Extract text from tables
            if doc.tables:
                for table_num, table in enumerate(doc.tables):
                    table_data = []
                    for row in table.rows:
                        row_data = [cell.text.strip() for cell in row.cells]
                        table_data.append(" | ".join(row_data))
                    
                    if table_data:
                        text_content.append(f"\n--- Table {table_num + 1} ---\n" + "\n".join(table_data))
            
            # Extract images if PIL is available
            if PILImage is not None:
                # Get all relationships that are images
                for rel_id, rel in doc.part.rels.items():
                    if "image" in rel.target_ref:
                        try:
                            # Get image data
                            image_blob = rel.target_part.blob
                            
                            # Convert to base64
                            img_base64 = base64.b64encode(image_blob).decode('utf-8')
                            
                            # Analyze image using common method
                            analysis = self._analyze_image_with_vision(
                                img_base64,
                                context="from DOCX"
                            )
                            
                            image_analyses.append(f"--- Image ---\n{analysis}")
                                
                        except Exception as e:
                            logger.warning(f"Failed to extract image: {e}")
            
            # Combine text and image analyses
            all_content = []
            
            if text_content:
                all_content.append("=== DOCX TEXT CONTENT ===\n" + "\n\n".join(text_content))
            
            if image_analyses:
                all_content.append("=== IMAGE ANALYSIS ===\n" + "\n\n".join(image_analyses))
            
            extracted_text = "\n\n".join(all_content)
            
            if not extracted_text.strip():
                logger.warning(f"No text or images extracted from DOCX: {file_path.name}")
                return f"[No content found in {file_path.name}]"
            
            return extracted_text
            
        except Exception as e:
            logger.error(f"Unexpected error extracting DOCX {file_path.name}: {e}")
            raise ValueError(f"Failed to extract DOCX: {e}")
    
    def _extract_image(self, file_path: Path) -> str:
        """Extract text from image using Azure OpenAI Vision"""
        try:
            import base64
        except ImportError:
            logger.error("base64 module not available")
            raise ValueError("base64 library not available")
        
        try:
            # Read image file and convert to base64
            with open(file_path, 'rb') as image_file:
                image_data = image_file.read()
                img_base64 = base64.b64encode(image_data).decode('utf-8')
            
            # Analyze image using common method
            analysis = self._analyze_image_with_vision(
                img_base64,
                context=f"standalone image {file_path.name}"
            )
            
            # Format output
            extracted_text = f"=== IMAGE ANALYSIS ===\n{analysis}"
            
            return extracted_text
            
        except Exception as e:
            logger.error(f"Failed to extract text from image {file_path.name}: {e}")
            raise ValueError(f"Failed to extract image text: {e}")
    
    def _extract_collection_data(self, content: str, metadata_dict: Dict) -> Dict[str, Dict]:
        """
        Extract structured data for 3 collections using Agent API
        
        Args:
            content: Combined content from all files
            metadata_dict: Original metadata dictionary
            
        Returns:
            Dict with keys 'invoice', 'vendors', 'skus' containing extracted data
        """
        try:
            from api.models import InvoiceDocument, VendorDocument, SKUDocument
        except ImportError:
            logger.warning("Could not import models from api.models, attempting alternative import")
            try:
                # Try alternative import paths
                import sys
                from pathlib import Path
                # Add api directory to path if not already there
                api_path = Path(__file__).parent.parent.parent / 'api'
                if str(api_path) not in sys.path:
                    sys.path.insert(0, str(api_path))
                from models import InvoiceDocument, VendorDocument, SKUDocument
            except ImportError as e:
                logger.error(f"Failed to import models: {e}")
                raise ImportError("Could not import Pydantic models. Ensure api/models.py is accessible.")
        
        # Get schema examples from models
        invoice_schema = InvoiceDocument.model_config.get('json_schema_extra', {}).get('example', {})
        vendor_schema = VendorDocument.model_config.get('json_schema_extra', {}).get('example', {})
        sku_schema = SKUDocument.model_config.get('json_schema_extra', {}).get('example', {})
        
        # Create prompt for agent to extract structured data with exact schemas
        prompt = f"""Analyze the following invoice/document content and extract structured information for three collections.

DOCUMENT CONTENT:
{content}

Extract and return JSON data matching these EXACT schemas:

1. INVOICE COLLECTION - Schema and Example:
{json.dumps(invoice_schema, indent=2)}

Required fields:
- invoice_id: DO NOT EXTRACT - will be auto-generated (string)
- content: Summarized invoice content suitable for semantic search (string)
- metadata: Object containing:
  * vendor_name: Vendor/supplier name (string)
  * vendor_id: DO NOT EXTRACT - will be auto-generated (string)
  * amount: Total amount (number)
  * currency: Currency code (string)
  * date: Invoice date (string)
  * po_number: Purchase order number (string)
  * language: Document language (string)
  [Do not provide attachments]

2. VENDOR COLLECTION - Schema and Example:
{json.dumps(vendor_schema, indent=2)}

Required fields:
- vendor_id: DO NOT EXTRACT - will be auto-generated (string)
- content: Vendor description including name, location, specialization (string)
- metadata: Object containing:
  * vendor_name: Vendor name (string)
  * country: Country of origin (string)
  * currency: Primary currency (string)
  * full_address: Complete vendor address (string)
  [Do not provide attachments]

3. SKU COLLECTION - Schema and Example:
{json.dumps(sku_schema, indent=2)}

Required fields:
- item_code: DO NOT EXTRACT - will be auto-generated (string)
- content: Item description suitable for semantic search (string)
- metadata: Object containing:
  * category: Product category (string)
  * uom: Unit of measure (string)
  * gst_rate: GST/tax rate (number)
  [Do not provide attachments]

Return your response in this EXACT JSON format:
{{
  "invoice": {{
    "content": "string",
    "metadata": {{
      "vendor_name": "string",
      "amount": 0.0,
      "currency": "string",
      "date": "string",
      "po_number": "string",
      "language": "string"
    }}
  }},
  "vendors": [
    {{
      "content": "string",
      "metadata": {{
        "vendor_name": "string",
        "country": "string",
        "currency": "string",
        "full_address": "string"
      }}
    }}
  ],
  "skus": [
    {{
      "content": "string",
      "metadata": {{
        "category": "string",
        "uom": "string",
        "gst_rate": 0
      }}
    }}
  ]
}}

IMPORTANT: 
- Return ONLY valid JSON, no markdown formatting, no code blocks, no additional text
- All metadata fields must be present even if empty/null
- Amounts and rates must be numbers, not strings
- Match the schema structure exactly
- vendors and skus MUST be arrays (even if only one item)
- Extract ALL vendors and ALL line items/SKUs from the invoice
- Provide ALL text content in English only - translate any non-English text to English
- All extracted data (content, metadata values, descriptions) must be in English
- DO NOT include invoice_id, vendor_id, or item_code in your response - these will be auto-generated"""

        # Retry logic: Try up to 3 times to get valid JSON from Agent API
        max_attempts = 3
        last_error = None
        
        for attempt in range(1, max_attempts + 1):
            try:
                if attempt > 1:
                    logger.warning(f"Retry attempt {attempt}/{max_attempts} for Agent API call...")
                
                # Call Azure OpenAI agent to extract structured data
                response = self.openai_client.chat(
                    prompt=prompt,
                    system_message="You are a data extraction expert. Extract structured data from invoices and return valid JSON only. Always respond in English, translating any non-English content to English.",
                    temperature=0.1,
                    max_tokens=2000
                )
                
                # Parse JSON response (json already imported at module level)
                collection_data = json.loads(response.strip())
                
                # Validate that all required keys are present
                required_keys = ['invoice', 'vendors', 'skus']
                for key in required_keys:
                    if key not in collection_data:
                        raise ValueError(f"Missing required key in response: {key}")
                
                # Validate that vendors and skus are arrays
                if not isinstance(collection_data['vendors'], list):
                    raise ValueError("vendors must be an array")
                if not isinstance(collection_data['skus'], list):
                    raise ValueError("skus must be an array")
                
                # Note: IDs will be generated AFTER acquiring locks to prevent race conditions
                logger.info(f"Extracted {len(collection_data['vendors'])} vendors and {len(collection_data['skus'])} SKUs")
                
                return collection_data
                
            except json.JSONDecodeError as e:
                last_error = e
                logger.error(f"Attempt {attempt}/{max_attempts}: Failed to parse Agent API response as JSON: {e}")
                if attempt >= max_attempts:
                    logger.error(f"All {max_attempts} attempts failed. Last response: {response[:500] if 'response' in locals() else 'No response'}")
                    raise ValueError(f"Invalid JSON response from Agent API after {max_attempts} attempts: {e}")
                    
            except ValueError as e:
                last_error = e
                logger.error(f"Attempt {attempt}/{max_attempts}: Validation error: {e}")
                if attempt >= max_attempts:
                    logger.error(f"All {max_attempts} attempts failed with validation errors")
                    raise
                    
            except Exception as e:
                last_error = e
                logger.error(f"Attempt {attempt}/{max_attempts}: Agent API call failed: {e}")
                if attempt >= max_attempts:
                    logger.error(f"All {max_attempts} attempts failed")
                    raise
        
        # This should never be reached, but just in case
        if last_error:
            raise last_error
        raise ValueError("Agent API extraction failed after all retry attempts")
    
    def create_embeddings(self, text: str) -> List[float]:
        """
        Create embeddings for text using Azure OpenAI
        
        Args:
            text: Text to create embeddings for
            
        Returns:
            List[float]: Embedding vector
        """
        try:
            # Truncate text if too long (max ~8000 tokens for text-embedding-3-large)
            max_chars = 30000  # Roughly 8000 tokens
            if len(text) > max_chars:
                logger.warning(f"Text too long ({len(text)} chars), truncating to {max_chars}")
                text = text[:max_chars]
            
            # Create embeddings using Azure OpenAI
            embeddings = self.openai_client.create_embeddings(text)
            return embeddings
        except Exception as e:
            logger.error(f"Failed to create embeddings: {e}")
            raise
    
    def store_in_vector_db(self, content: str, metadata: Dict, doc_id: str, embeddings: List[float]) -> None:
        """
        Store document in vector database with embeddings
        
        Args:
            content: Text content to store
            metadata: Document metadata
            doc_id: Unique document ID
            embeddings: Pre-computed embedding vector from Azure OpenAI
        """
        try:
            # Store in vector database with your Azure embeddings
            self.vector_store.add_invoice(
                invoice_id=doc_id,
                content=content,
                metadata=metadata,
                embedding=embeddings  # Pass Azure embeddings
            )
            
        except Exception as e:
            logger.error(f"Failed to store in vector database: {e}")
            raise
    
    def load_invoice(self, meta_file: Path, attachment_paths: List[Path], metadata_dict: Dict) -> str:
        """
        Main method: Extract content, create embeddings, and store in vector database
        
        Collection locks prevent race conditions during parallel processing.
        
        Args:
            meta_file: Path to metadata file
            attachment_paths: List of attachment file paths
            metadata_dict: Metadata dictionary from InvoiceMetadata
            
        Returns:
            str: "OK" if successful, error message if failed
        """
        try:
            # Step 1: Extract content from all files
            all_content = []
            
            # Extract from metadata file
            try:
                with open(meta_file, 'r', encoding='utf-8') as f:
                    meta_content = json.load(f)
                # Remove attachments field before sending to agent (not part of schema)
                meta_content.pop('attachments', None)
                all_content.append(f"Metadata: {json.dumps(meta_content, indent=2)}")
            except Exception as e:
                logger.warning(f"Could not extract metadata content: {e}")
            
            # Extract from attachments
            for file_path in attachment_paths:
                try:
                    content = self.extract_content(file_path)
                    all_content.append(f"\n--- {file_path.name} ---\n{content}")
                except Exception as e:
                    error_msg = f"Failed to extract content from {file_path.name}: {e}"
                    logger.error(error_msg)
                    return error_msg
            
            # Combine all content
            combined_content = "\n\n".join(all_content)
            
            # Step 2: Extract structured data using Agent API
            try:
                collection_data = self._extract_collection_data(combined_content, metadata_dict)
            except Exception as e:
                error_msg = f"Failed to extract collection data: {e}"
                logger.error(error_msg)
                return error_msg
            
            # Step 3: Create embeddings and store in vector database
            try:
                # First, generate ALL embeddings in parallel (slow part - no locks needed)
                logger.info("Generating embeddings for invoice, vendors, and SKUs...")
                
                # Prepare invoice data and embeddings
                invoice_data = collection_data['invoice']
                invoice_content = invoice_data.get('content', combined_content).strip()
                
                if not invoice_content:
                    invoice_id = invoice_data.get('invoice_id', 'Unknown')
                    invoice_content = f"Invoice {invoice_id}"
                    logger.warning(f"Empty invoice content for {invoice_id}, using fallback")
                
                invoice_embeddings = self.create_embeddings(invoice_content)
                
                # Prepare vendor data and embeddings
                vendor_data_list = []
                for vendor_data in collection_data['vendors']:
                    vendor_content = vendor_data.get('content', '').strip()
                    
                    if not vendor_content:
                        vendor_metadata_raw = vendor_data.get('metadata', {})
                        vendor_name = vendor_metadata_raw.get('vendor_name', 'Unknown Vendor')
                        vendor_content = f"Vendor: {vendor_name}"
                        logger.warning(f"Empty vendor content for {vendor_data.get('vendor_id')}, using fallback")
                    
                    vendor_embeddings = self.create_embeddings(vendor_content)
                    vendor_data_list.append({
                        'data': vendor_data,
                        'content': vendor_content,
                        'embeddings': vendor_embeddings
                    })
                
                # Prepare SKU data and embeddings
                sku_data_list = []
                for sku_data in collection_data['skus']:
                    sku_content = sku_data.get('content', '').strip()
                    
                    if not sku_content:
                        sku_metadata_raw = sku_data.get('metadata', {})
                        sku_desc = sku_metadata_raw.get('description', 'Unknown SKU')
                        sku_content = f"SKU: {sku_desc}"
                        logger.warning(f"Empty SKU content for {sku_data.get('item_code')}, using fallback")
                    
                    sku_embeddings = self.create_embeddings(sku_content)
                    sku_data_list.append({
                        'data': sku_data,
                        'content': sku_content,
                        'embeddings': sku_embeddings
                    })
                
                logger.info(f"Embeddings generated. Now processing inserts with individual locks...")
                
                # === INVOICE: Lock → Generate ID → Insert → Unlock ===
                try:
                    requests.post(f"{self.api_base_url}/vector/lock/invoices_stage", timeout=60)
                    logger.info("Invoice collection locked")
                except Exception as lock_error:
                    logger.warning(f"Failed to lock invoice collection: {lock_error}")
                
                try:
                    # Generate invoice ID
                    invoice_id = self._get_next_invoice_id()
                    logger.info(f"Generated invoice ID: {invoice_id}")
                    
                    # Inject ID into invoice data
                    invoice_data['invoice_id'] = invoice_id
                    
                    # Prepend invoice ID to content for better search
                    invoice_content_with_id = f"Invoice ID: {invoice_id}. {invoice_content}"
                    
                    # Store invoice data
                    invoice_metadata = invoice_data.get('metadata', {})
                    # Remove attachments field if present (not part of schema)
                    invoice_metadata.pop('attachments', None)
                    invoice_metadata['meta_file'] = meta_file.name
                    invoice_metadata['processed_timestamp'] = datetime.now().isoformat()
                    invoice_metadata['review_status'] = 'pending'
                    invoice_metadata['reviewed_by'] = None
                    invoice_metadata['reviewed_at'] = None
                    invoice_metadata['review_notes'] = None
                    invoice_metadata['original_values'] = json.dumps({})
                    
                    self.vector_store.add_invoice_stage(
                        invoice_id=invoice_data['invoice_id'],
                        content=invoice_content_with_id,
                        metadata=invoice_metadata,
                        embedding=invoice_embeddings
                    )
                    logger.info(f"✓ Inserted invoice {invoice_id}")
                finally:
                    try:
                        requests.post(f"{self.api_base_url}/vector/unlock/invoices_stage", timeout=10)
                    except Exception as unlock_error:
                        logger.warning(f"Failed to unlock invoice collection: {unlock_error}")
                
                # === VENDORS: Lock → Get Starting ID → Generate Rest Locally → Insert All → Unlock ===
                generated_vendor_ids = []
                if vendor_data_list:
                    try:
                        requests.post(f"{self.api_base_url}/vector/lock/vendors_stage", timeout=60)
                        logger.info("Vendor collection locked")
                    except Exception as lock_error:
                        logger.warning(f"Failed to lock vendor collection: {lock_error}")
                    
                    try:
                        # Get starting vendor ID from API (single call)
                        base_vendor_response = requests.get(f"{self.api_base_url}/vector/next-id/vendor", timeout=60)
                        base_vendor_response.raise_for_status()
                        base_vendor_id_str = base_vendor_response.json()['id']
                        
                        # Extract numeric part and generate rest locally
                        base_vendor_num = int(base_vendor_id_str.replace('VND-', ''))
                        generated_vendor_ids = [f"VND-{base_vendor_num + i:07d}" for i in range(len(vendor_data_list))]
                        logger.info(f"Generated {len(generated_vendor_ids)} vendor IDs locally from {base_vendor_id_str}")
                        
                        # Update invoice metadata with first vendor ID
                        if generated_vendor_ids and 'metadata' in invoice_data:
                            invoice_data['metadata']['vendor_id'] = generated_vendor_ids[0]
                        
                        # Insert all vendors
                        vendor_ids = []
                        for idx, vendor_item in enumerate(vendor_data_list):
                            vendor_data = vendor_item['data']
                            vendor_content = vendor_item['content']
                            vendor_embeddings = vendor_item['embeddings']
                            
                            # Inject generated vendor ID
                            vendor_id = generated_vendor_ids[idx]
                            vendor_data['vendor_id'] = vendor_id
                            
                            # Prepend vendor ID to content for better search
                            vendor_content_with_id = f"Vendor ID: {vendor_id}. {vendor_content}"
                            
                            # Add metadata
                            vendor_metadata = vendor_data.get('metadata', {})
                            # Remove attachments field if present (not part of schema)
                            vendor_metadata.pop('attachments', None)
                            vendor_metadata['src_invoice_id'] = invoice_data['invoice_id']
                            vendor_metadata['review_status'] = 'pending'
                            vendor_metadata['reviewed_by'] = None
                            vendor_metadata['reviewed_at'] = None
                            vendor_metadata['review_notes'] = None
                            vendor_metadata['original_values'] = json.dumps({})
                            
                            self.vector_store.add_vendor_stage(
                                vendor_id=vendor_data['vendor_id'],
                                content=vendor_content_with_id,
                                metadata=vendor_metadata,
                                embedding=vendor_embeddings
                            )
                            vendor_ids.append(vendor_data['vendor_id'])
                        
                        logger.info(f"✓ Inserted {len(vendor_ids)} vendors")
                    finally:
                        try:
                            requests.post(f"{self.api_base_url}/vector/unlock/vendors_stage", timeout=10)
                        except Exception as unlock_error:
                            logger.warning(f"Failed to unlock vendor collection: {unlock_error}")
                
                # === SKUs: Lock → Get Starting ID → Generate Rest Locally → Insert All → Unlock ===
                generated_sku_ids = []
                if sku_data_list:
                    try:
                        requests.post(f"{self.api_base_url}/vector/lock/skus_stage", timeout=60)
                        logger.info("SKU collection locked")
                    except Exception as lock_error:
                        logger.warning(f"Failed to lock SKU collection: {lock_error}")
                    
                    try:
                        # Get starting SKU ID from API (single call)
                        base_sku_response = requests.get(f"{self.api_base_url}/vector/next-id/sku", timeout=60)
                        base_sku_response.raise_for_status()
                        base_sku_id_str = base_sku_response.json()['id']
                        
                        # Extract numeric part and generate rest locally
                        base_sku_num = int(base_sku_id_str.replace('SKU-', ''))
                        generated_sku_ids = [f"SKU-{base_sku_num + i:07d}" for i in range(len(sku_data_list))]
                        logger.info(f"Generated {len(generated_sku_ids)} SKU IDs locally from {base_sku_id_str}")
                        
                        # Insert all SKUs
                        sku_ids = []
                        for idx, sku_item in enumerate(sku_data_list):
                            sku_data = sku_item['data']
                            sku_content = sku_item['content']
                            sku_embeddings = sku_item['embeddings']
                            
                            # Inject generated SKU ID
                            sku_id = generated_sku_ids[idx]
                            sku_data['item_code'] = sku_id
                            
                            # Prepend SKU ID to content for better search
                            sku_content_with_id = f"SKU ID: {sku_id}. {sku_content}"
                            
                            # Add metadata
                            sku_metadata = sku_data.get('metadata', {})
                            # Remove attachments field if present (not part of schema)
                            sku_metadata.pop('attachments', None)
                            sku_metadata['src_invoice_id'] = invoice_data['invoice_id']
                            sku_metadata['review_status'] = 'pending'
                            sku_metadata['reviewed_by'] = None
                            sku_metadata['reviewed_at'] = None
                            sku_metadata['review_notes'] = None
                            sku_metadata['original_values'] = json.dumps({})
                            
                            self.vector_store.add_sku_stage(
                                item_code=sku_data['item_code'],
                                content=sku_content_with_id,
                                metadata=sku_metadata,
                                embedding=sku_embeddings
                            )
                            sku_ids.append(sku_data['item_code'])
                        
                        logger.info(f"✓ Inserted {len(sku_ids)} SKUs")
                    finally:
                        try:
                            requests.post(f"{self.api_base_url}/vector/unlock/skus_stage", timeout=10)
                        except Exception as unlock_error:
                            logger.warning(f"Failed to unlock SKU collection: {unlock_error}")
                
                # Final summary
                logger.info(f"✓ Completed: Invoice {invoice_id}, {len(generated_vendor_ids)} vendors, {len(generated_sku_ids)} SKUs")
                # Final summary
                logger.info(f"✓ Completed: Invoice {invoice_id}, {len(generated_vendor_ids)} vendors, {len(generated_sku_ids)} SKUs")
                
                return "OK"
                
            except Exception as e:
                error_msg = f"Failed to store in vector database: {e}"
                logger.error(error_msg)
                return error_msg
            
        except Exception as e:
            error_msg = f"Unexpected error during vector DB loading: {e}"
            logger.error(error_msg)
            return error_msg